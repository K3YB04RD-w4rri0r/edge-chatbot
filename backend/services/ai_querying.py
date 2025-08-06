import os
import logging
import base64
from typing import List, Dict, Any, Optional
from datetime import datetime
import asyncio
from openai import AsyncAzureOpenAI  # Use Azure OpenAI client
import google.generativeai as genai
import tiktoken
import tempfile
from pathlib import Path

from models.conversations_model import ModelChoice, ModelInstructions, Conversation
from models.messages_model import Message, MessageRole
from models.attachments_model import AttachmentType

from shared_variables import settings

logger = logging.getLogger(__name__)

import os
import logging
import base64
from typing import List, Dict, Any, Optional
from datetime import datetime
import asyncio
from openai import AsyncAzureOpenAI
import google.generativeai as genai
import tiktoken
import tempfile
from pathlib import Path
import mimetypes
import io

# Document processing libraries
import PyPDF2
import docx
import openpyxl
import csv
import json



from models.conversations_model import ModelChoice, ModelInstructions, Conversation
from models.messages_model import Message, MessageRole
from models.attachments_model import AttachmentType

from shared_variables import settings

logger = logging.getLogger(__name__)

# Token encoder for OpenAI models
encoding = tiktoken.encoding_for_model("gpt-4")

# Maximum tokens for document context (adjust based on your needs)
MAX_DOCUMENT_TOKENS = 8000  # Reserve space for conversation history and response


class DocumentExtractor:
    """Extract text from various document types"""
    
    @staticmethod
    async def extract_text(file_content: bytes, filename: str, content_type: str = None) -> str:
        """
        Extract text from document based on file type.
        Returns extracted text or empty string if extraction fails.
        """
        try:
            # Determine file type
            file_ext = Path(filename).suffix.lower()
            if not content_type:
                content_type, _ = mimetypes.guess_type(filename)
            
            # Route to appropriate extractor
            if file_ext == '.pdf' or content_type == 'application/pdf':
                return await DocumentExtractor._extract_pdf(file_content)
            
            elif file_ext in ['.docx', '.doc'] or content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                return await DocumentExtractor._extract_docx(file_content)
            
            elif file_ext in ['.xlsx', '.xls'] or content_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
                return await DocumentExtractor._extract_excel(file_content)
            
            elif file_ext == '.csv' or content_type == 'text/csv':
                return await DocumentExtractor._extract_csv(file_content)
            
            elif file_ext == '.json' or content_type == 'application/json':
                return await DocumentExtractor._extract_json(file_content)
            
            elif file_ext in ['.txt', '.md', '.log'] or (content_type and content_type.startswith('text/')):
                return await DocumentExtractor._extract_text_file(file_content)
            
            elif file_ext in ['.py', '.js', '.java', '.cpp', '.c', '.html', '.css', '.xml']:
                # Code files
                return await DocumentExtractor._extract_text_file(file_content)
            
            else:
                logger.warning(f"Unsupported file type: {filename} ({content_type})")
                return f"[Unable to extract text from {filename}]"
                
        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {e}")
            return f"[Error extracting text from {filename}: {str(e)}]"
    
    @staticmethod
    async def _extract_pdf(file_content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"[Page {page_num}]\n{page_text}")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return "[Failed to extract PDF content]"
    
    @staticmethod
    async def _extract_docx(file_content: bytes) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return "[Failed to extract Word document content]"
    
    @staticmethod
    async def _extract_excel(file_content: bytes) -> str:
        """Extract text from Excel file"""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True, data_only=True)
            sheets_text = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []
                
                for row in sheet.iter_rows(values_only=True):
                    # Filter out empty rows
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(val.strip() for val in row_values):
                        sheet_data.append(" | ".join(row_values))
                
                if sheet_data:
                    sheets_text.append(f"[Sheet: {sheet_name}]\n" + "\n".join(sheet_data[:100]))  # Limit rows
            
            return "\n\n".join(sheets_text)
        except Exception as e:
            logger.error(f"Excel extraction error: {e}")
            return "[Failed to extract Excel content]"
    
    @staticmethod
    async def _extract_csv(file_content: bytes) -> str:
        """Extract text from CSV file"""
        try:
            text = file_content.decode('utf-8', errors='ignore')
            reader = csv.reader(io.StringIO(text))
            rows = []
            
            for i, row in enumerate(reader):
                if i >= 100:  # Limit to first 100 rows
                    rows.append("[... truncated ...]")
                    break
                rows.append(" | ".join(row))
            
            return "\n".join(rows)
        except Exception as e:
            logger.error(f"CSV extraction error: {e}")
            return "[Failed to extract CSV content]"
    
    @staticmethod
    async def _extract_json(file_content: bytes) -> str:
        """Extract text from JSON file"""
        try:
            data = json.loads(file_content.decode('utf-8', errors='ignore'))
            # Pretty print with limited depth
            return json.dumps(data, indent=2, ensure_ascii=False)[:10000]  # Limit size
        except Exception as e:
            logger.error(f"JSON extraction error: {e}")
            return "[Failed to extract JSON content]"
    
    @staticmethod
    async def _extract_text_file(file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all fail, use ignore errors
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            return "[Failed to extract text content]"


class TokenManager:
    """Manage token limits for document context"""
    
    @staticmethod
    def truncate_to_token_limit(text: str, max_tokens: int) -> str:
        """Truncate text to fit within token limit"""
        try:
            tokens = encoding.encode(text)
            if len(tokens) <= max_tokens:
                return text
            
            # Truncate and add indicator
            truncated_tokens = tokens[:max_tokens - 20]  # Reserve space for truncation message
            truncated_text = encoding.decode(truncated_tokens)
            return f"{truncated_text}\n\n[... content truncated due to length ...]"
        except Exception as e:
            logger.error(f"Token truncation error: {e}")
            # Fallback to character-based truncation
            char_limit = max_tokens * 4  # Approximate
            if len(text) > char_limit:
                return f"{text[:char_limit]}\n\n[... content truncated due to length ...]"
            return text


# Update the Azure OpenAI chat completion function
async def _handle_azure_chat_completion_with_documents(
    messages: List[Message],
    new_message: str,
    active_attachments: List[Dict[str, Any]],
    model_choice: str,
    model_instructions: str
) -> str:
    """Handle Azure OpenAI chat completion with document extraction"""
    
    formatted_messages = []
    
    # Add system message
    formatted_messages.append({
        "role": "system",
        "content": model_instructions
    })
    
    # Add conversation history (limited to preserve token space)
    history_messages = messages[-10:]  # Reduce history to save tokens for documents
    for message in history_messages:
        formatted_messages.append({
            "role": message.role,
            "content": message.content
        })
    
    # Process attachments
    document_contexts = []
    image_contents = []
    
    for attachment in active_attachments:
        if attachment.get("type") == AttachmentType.IMAGE.value:
            # Handle images
            if "file_content" in attachment:
                base64_image = base64.b64encode(
                    attachment["file_content"]
                ).decode('utf-8')
                
                image_contents.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{attachment.get('content_type', 'image/jpeg')};base64,{base64_image}",
                        "detail": "high"
                    }
                })
        else:
            # Extract text from documents
            if "file_content" in attachment:
                extracted_text = await DocumentExtractor.extract_text(
                    attachment["file_content"],
                    attachment["filename"],
                    attachment.get("content_type")
                )
                
                if extracted_text and extracted_text.strip():
                    document_contexts.append(
                        f"### Document: {attachment['filename']}\n{extracted_text}"
                    )
    
    # Combine document contexts with token management
    user_content_parts = []
    
    if document_contexts:
        # Calculate available tokens for documents
        tokens_per_doc = MAX_DOCUMENT_TOKENS // len(document_contexts)
        
        truncated_contexts = []
        for context in document_contexts:
            truncated = TokenManager.truncate_to_token_limit(context, tokens_per_doc)
            truncated_contexts.append(truncated)
        
        combined_context = "\n\n---\n\n".join(truncated_contexts)
        
        # Build message with document context
        if image_contents:
            # Mixed content with images and documents
            user_content_parts = [
                {"type": "text", "text": f"## Attached Documents:\n{combined_context}\n\n## User Message:\n{new_message}"}
            ] + image_contents
        else:
            # Text-only with documents
            user_content_parts = f"## Attached Documents:\n{combined_context}\n\n## User Message:\n{new_message}"
    elif image_contents:
        # Images only
        user_content_parts = [
            {"type": "text", "text": new_message}
        ] + image_contents
    else:
        # No attachments
        user_content_parts = new_message
    
    # Add user message
    formatted_messages.append({
        "role": "user",
        "content": user_content_parts
    })
    
    # Log token usage estimate
    try:
        total_text = " ".join(
            msg["content"] if isinstance(msg["content"], str) else 
            msg["content"][0]["text"] if isinstance(msg["content"], list) else ""
            for msg in formatted_messages
        )
        estimated_tokens = len(encoding.encode(total_text))
        logger.info(f"Estimated input tokens: {estimated_tokens}")
    except Exception as e:
        logger.warning(f"Could not estimate tokens: {e}")
    
    # Azure deployment name mapping
    deployment_mapping = {
        ModelChoice.GPT_4_1_NANO.value: "gpt-4.1-nano",
        ModelChoice.GPT_4_1.value: "gpt-4.1"
    }
    
    # Make API call
    try:
        response = await azure_openai_client.chat.completions.create(
            model=deployment_mapping.get(model_choice, "gpt-4.1"),
            messages=formatted_messages,
            max_tokens=4096,
            temperature=0.7
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        if "context_length_exceeded" in str(e).lower():
            logger.error("Context length exceeded. Retrying with reduced context...")
            # Retry with more aggressive truncation
            # You could implement a retry strategy here
            raise Exception("Document content too large. Please upload smaller documents or fewer files.")
        raise


azure_openai_client = AsyncAzureOpenAI(
    api_key=settings.azure_openai_key,  
    api_version="2025-01-01-preview",  
    azure_endpoint=settings.azure_endpoint 
)

# For Gemini
genai.configure(api_key=settings.google_key)

# Token encoder for OpenAI models
encoding = tiktoken.encoding_for_model("gpt-4")


class FileUploadManager:
    """Manages file uploads for both Azure OpenAI and Gemini"""
    
    def __init__(self):
        self.azure_files = {}  # Cache Azure file IDs
        self.gemini_files = {}  # Cache Gemini file objects
    
    async def upload_to_azure(self, file_content: bytes, filename: str) -> str:
        """Upload file to Azure OpenAI and return file ID"""
        try:
            import io
            file_object = io.BytesIO(file_content)
            file_object.name = filename
            
            # Note: Azure OpenAI file upload may have different requirements
            # Check if your Azure deployment supports file uploads
            file_response = await azure_openai_client.files.create(
                file=file_object,
                purpose="assistants"
            )
            
            logger.info(f"Uploaded {filename} to Azure OpenAI: {file_response.id}")
            return file_response.id
            
        except Exception as e:
            logger.error(f"Failed to upload to Azure OpenAI: {e}")
            raise
    
    def upload_to_gemini(self, file_content: bytes, filename: str) -> Any:
        """Upload file to Gemini and return file object"""
        try:
            # Gemini requires a file path, so create temp file
            suffix = Path(filename).suffix
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            
            # Upload to Gemini
            gemini_file = genai.upload_file(
                path=tmp_path,
                display_name=filename
            )
            
            # Clean up temp file
            os.unlink(tmp_path)
            
            logger.info(f"Uploaded {filename} to Gemini: {gemini_file.uri}")
            return gemini_file
            
        except Exception as e:
            logger.error(f"Failed to upload to Gemini: {e}")
            raise
    
    async def cleanup_azure_files(self, file_ids: List[str]):
        """Delete Azure OpenAI files after use"""
        for file_id in file_ids:
            try:
                await azure_openai_client.files.delete(file_id)
                logger.info(f"Deleted Azure file: {file_id}")
            except Exception as e:
                logger.error(f"Failed to delete file {file_id}: {e}")


# Global file manager
file_manager = FileUploadManager()


async def get_ai_response(
    conversation: Conversation,
    messages: List[Message],
    new_message: str,
    active_attachments: List[Dict[str, Any]],
    model_choice: str,
    model_instructions: str
) -> str:
    """
    Generate AI response with direct file upload support.
    Supports both Azure OpenAI and Gemini models.
    """
    try:
        # Check if it's an Azure OpenAI model (your custom deployment names)
        if model_choice in [ModelChoice.GPT_4_1_NANO.value, ModelChoice.GPT_4_1.value]:
            return await _handle_azure_openai_request(
                conversation, messages, new_message, 
                active_attachments, model_choice, model_instructions
            )
        elif model_choice in [ModelChoice.GEMINI_2_FLASH_EXP]:
            return await _handle_gemini_request(
                messages, new_message, active_attachments, model_instructions, model_choice
            )
        else:
            raise ValueError(f"Unsupported model: {model_choice}")
            
    except Exception as e:
        logger.error(f"Failed to generate AI response: {e}")
        raise


async def _handle_azure_openai_request(
    conversation: Conversation,
    messages: List[Message],
    new_message: str,
    active_attachments: List[Dict[str, Any]],
    model_choice: str,
    model_instructions: str
) -> str:
    """Handle Azure OpenAI requests with document text extraction"""
    
    # Use the new function that handles document extraction
    return await _handle_azure_chat_completion_with_documents(
        messages, new_message, active_attachments,
        model_choice, model_instructions
    )


async def _handle_azure_chat_completion(
    messages: List[Message],
    new_message: str,
    active_attachments: List[Dict[str, Any]],
    model_choice: str,
    model_instructions: str,
    has_images: bool
) -> str:
    """Handle Azure OpenAI chat completion with optional images"""
    
    formatted_messages = []
    
    # Add system message
    formatted_messages.append({
        "role": "system",
        "content": model_instructions
    })
    
    # Add conversation history
    for message in messages[-20:]:  # Last 20 messages
        formatted_messages.append({
            "role": message.role,
            "content": message.content
        })
    
    # Build user message
    if has_images:
        user_content = [{"type": "text", "text": new_message}]
        
        for attachment in active_attachments:
            if attachment.get("type") == AttachmentType.IMAGE.value:
                if "file_content" in attachment:
                    # Convert to base64
                    base64_image = base64.b64encode(
                        attachment["file_content"]
                    ).decode('utf-8')
                    
                    user_content.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{attachment.get('content_type', 'image/jpeg')};base64,{base64_image}",
                            "detail": "high"  # or "low" for faster processing
                        }
                    })
        
        formatted_messages.append({
            "role": "user",
            "content": user_content
        })
    else:
        # Text-only message
        formatted_messages.append({
            "role": "user",
            "content": new_message
        })
    
    # Azure deployment name mapping
    # These should match your actual Azure deployment names
    deployment_mapping = {
        ModelChoice.GPT_4_1_NANO.value: "gpt-4.1-nano",  # Your Azure deployment name
        ModelChoice.GPT_4_1.value: "gpt-4.1"             # Your Azure deployment name
    }
    
    # Use Azure OpenAI chat completion
    response = await azure_openai_client.chat.completions.create(
        model=deployment_mapping.get(model_choice, "gpt-4.1"),  # This is your deployment name
        messages=formatted_messages,
        max_tokens=4096,
        temperature=0.7
    )
    
    return response.choices[0].message.content


async def _handle_gemini_request(
    messages: List[Message],
    new_message: str,
    active_attachments: List[Dict[str, Any]],
    model_instructions: str,
    model_choice,
) -> str:
    """Handle Gemini requests with native file support"""
    
    try:
        model_mapping = {ModelChoice.GEMINI_2_FLASH_EXP.value : "gemini-2.0-flash-exp"}

        model = genai.GenerativeModel(model_mapping.get(model_choice, "gemini-2.0-flash-exp" ))
        
        # Build conversation history
        chat_history = []
        for msg in messages[-20:]:  # Last 20 messages
            role = "user" if msg.role == MessageRole.USER.value else "model"
            chat_history.append({"role": role, "parts": [msg.content]})

        # Upload files and prepare content
        content_parts = []
        
        # Add files first
        for attachment in active_attachments:
            if "file_content" in attachment:
                # Upload to Gemini
                gemini_file = file_manager.upload_to_gemini(
                    attachment["file_content"],
                    attachment["filename"]
                )
                content_parts.append(gemini_file)
        
        # Add the user message
        content_parts.append(new_message)
        
        # Start chat with history
        chat = model.start_chat(history=chat_history)
        
        # Send message with files
        response = chat.send_message(content_parts)
        
        return response.text
        
    except Exception as e:
        logger.error(f"Gemini API error: {e}")
        raise


# Token counting helper
def count_tokens_approximate(messages: List[Message]) -> int:
    """Approximate token count for cost estimation"""
    try:
        full_text = " ".join(msg.content for msg in messages)
        tokens = encoding.encode(full_text)
        return len(tokens)
    except Exception:
        # Fallback: ~4 characters per token
        char_count = sum(len(msg.content) for msg in messages)
        return char_count // 4